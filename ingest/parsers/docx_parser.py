from __future__ import annotations

from pathlib import Path

from ingest.parsers.base import ParsedDocument, ParsedSection


class DOCXParser:
    """Parse DOCX files using python-docx."""

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(path))
        sections: list[ParsedSection] = []
        current_heading = ""
        current_texts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # Flush current section
                if current_texts:
                    sections.append(ParsedSection(
                        text="\n".join(current_texts),
                        heading=current_heading,
                    ))
                    current_texts = []
                current_heading = text
            else:
                current_texts.append(text)

        # Flush remaining
        if current_texts:
            sections.append(ParsedSection(
                text="\n".join(current_texts),
                heading=current_heading,
            ))

        if not sections:
            sections.append(ParsedSection(text="", heading=""))

        return ParsedDocument(
            title=path.stem,
            source_path=str(path),
            sections=sections,
        )
